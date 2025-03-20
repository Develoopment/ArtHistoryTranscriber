from docx import Document
from docx.shared import Inches
import os
import shutil

def read_md_file(file_path):

    with open(file_path, "r", encoding="utf-8") as file:
        #this object stores all the info in easy to read format
        insert_obj = {}

        #get the file name (which is the art title in md files)
        file_name = os.path.basename(file_path)
        art_title = os.path.splitext(file_name)[0]
        insert_obj["Title"] = art_title

        content = file.read().split("##")

        #converting the pasted path to the correct format needed to resolve the path to the image (second brain > Files >)
        #If I forget to put the img in the file
        try:
            pasted_img = content[0]
            img_name = pasted_img[3:pasted_img.index("]")]
            img_path = "C:/Users/neerr/Desktop/second brain/Files/" + img_name

            insert_obj["Img_Path"] = img_path
        except:
            print("You forgot to put the image for: " + art_title)
            insert_obj["Img_Path"] = "C:/Users/neerr/Desktop/second brain/Files/placeholder.jpg"

        ##getting contextual information
        con_info = content[3][12:]
        insert_obj["Contextual"] = con_info
        
        ##getting visual information
        vis_info = content[2][7:]
        insert_obj["Visual"] = vis_info
        
        ## getting art identification info
        id_info = content[1]

        #cleaning any extra newlines
        clean_set = set(id_info.split("\n")) #cnverting set removes all duplicate instances of "\n"
        clean_set.remove(" Id Info") #remove all instances of "" from the set
        clean_set.remove("")
        cleaned_info = list(clean_set)

        id_info_dict = {}

        for info in cleaned_info:
            key, value = info.split(":")
            key = key.strip()
            value = value.strip()

            id_info_dict[key] = value

        insert_obj["id_info"] = id_info_dict

        #returning completed object
        return(insert_obj)

        file.close()

def write_template(art_info):
    doc = Document('template.docx')
    
    # for paragraph in doc.paragraphs:
    #     print(paragraph.text) 

    #replacing title
    doc.paragraphs[0].text = art_info["Title"]
    
    #adding the art id info
    run_ref = doc.paragraphs[1]
    for item in art_info["id_info"]:
        run_ref.add_run(item + ": " + art_info["id_info"][item] + "\n")

    doc.add_picture = doc.add_picture(art_info["Img_Path"], width=Inches(6.52), height=Inches(4))
    #doc.add_picture(art_info["Img_Path"], width=Inches(3), height=Inches(4))

    #adding table
    table = doc.add_table(rows=2, cols=2, style="Table Grid")

    header_cells = table.rows[0].cells
    content_cells = table.rows[1].cells

    header_cells[0].text = "Visual"
    header_cells[1].text = "Contextual"
    
    content_cells[0].text = art_info["Visual"]
    content_cells[1].text = art_info["Contextual"]

    doc.save(art_info["Title"] + ".docx")
    doc_path = "C:/Users/neerr/Desktop/Code/Art History Transcribing/" + art_info["Title"] + ".docx"
    shutil.move(doc_path, "C:/Users/neerr/Desktop/Code/Art History Transcribing/Documents")
    print("Generated: " + art_info["Title"])

###MAIN LOOP
input("run program")
md_base_path = "C:/Users/neerr/Desktop/second brain/Art History/"
dir_list = os.listdir(md_base_path)

for file in dir_list:
    art_obj = read_md_file(md_base_path + file)
    write_template(art_obj)
